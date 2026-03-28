"""Resume manager — parse and generate .docx resumes.

Supported input formats for parse_resume():
  - .docx  — heuristic paragraph/heading extractor using python-docx
  - .pdf   — text-layer extraction via pypdf (does not work for scanned PDFs)
  - .md    — Markdown heading-based section extractor

The upload workflow is heuristic-based (no LLM). Sections are inferred by
keyword matching on heading text (see _map_resume_to_profile in
dashboard/pages/profile.py). For richer extraction, an LLM post-processing
step could be added in future.
"""

from __future__ import annotations

import logging
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from docx import Document
from docx.text.paragraph import Paragraph

logger = logging.getLogger(__name__)


@dataclass
class ResumeSection:
    """A logical section of the resume (e.g., Experience, Skills)."""

    heading: str
    items: list[str] = field(default_factory=list)


@dataclass
class ResumeData:
    """Structured representation of a parsed resume."""

    name: str = ""
    contact_info: str = ""
    summary: str = ""
    sections: list[ResumeSection] = field(default_factory=list)

    def get_section(self, heading: str) -> ResumeSection | None:
        """Find a section by heading (case-insensitive)."""
        for section in self.sections:
            if section.heading.lower() == heading.lower():
                return section
        return None

    def to_dict(self) -> dict[str, Any]:
        return {
            "name": self.name,
            "contact_info": self.contact_info,
            "summary": self.summary,
            "sections": [
                {"heading": s.heading, "items": s.items} for s in self.sections
            ],
        }

    @classmethod
    def from_dict(cls, data: dict[str, Any]) -> ResumeData:
        sections = [
            ResumeSection(heading=s["heading"], items=s["items"])
            for s in data.get("sections", [])
        ]
        return cls(
            name=data.get("name", ""),
            contact_info=data.get("contact_info", ""),
            summary=data.get("summary", ""),
            sections=sections,
        )


def _is_heading(paragraph: Paragraph) -> bool:
    """Check if a paragraph is a heading (by style or bold formatting)."""
    style_name = (paragraph.style.name or "").lower()
    if "heading" in style_name:
        return True
    # Check if entire paragraph is bold and short (likely a section header)
    if paragraph.runs and all(r.bold for r in paragraph.runs if r.text.strip()):
        text = paragraph.text.strip()
        if text and len(text) < 80:
            return True
    return False


def parse_resume(path: Path) -> ResumeData:
    """Parse a resume file into structured ResumeData.

    Dispatches to the appropriate parser based on file extension:
    - .docx  — heuristic paragraph/heading extractor
    - .pdf   — text-layer extraction (text-based PDFs only)
    - .md    — Markdown heading-based extractor
    """
    ext = path.suffix.lower()
    if ext == ".docx":
        return _parse_resume_docx(path)
    elif ext == ".pdf":
        return _parse_resume_pdf(path)
    elif ext == ".md":
        return _parse_resume_md(path)
    else:
        raise ValueError(f"Unsupported resume format: {ext!r}. Use .docx, .pdf, or .md")


def _parse_resume_docx(path: Path) -> ResumeData:
    """Parse a .docx resume into structured ResumeData.

    Heuristic approach:
    - First non-empty paragraph = name
    - Heading-style paragraphs start new sections
    - Everything else is items within the current section
    """
    doc = Document(str(path))
    data = ResumeData()
    current_section: ResumeSection | None = None
    found_name = False

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # First non-empty text is the name
        if not found_name:
            data.name = text
            found_name = True
            continue

        # Second before any heading is contact info
        if not data.contact_info and current_section is None and not _is_heading(para):
            data.contact_info = text
            continue

        if _is_heading(para):
            current_section = ResumeSection(heading=text)
            data.sections.append(current_section)
        elif current_section is not None:
            current_section.items.append(text)
        else:
            # Text before any section heading — treat as summary
            if data.summary:
                data.summary += "\n" + text
            else:
                data.summary = text

    logger.info("Parsed resume (docx): %s, %d sections", data.name, len(data.sections))
    return data


def _parse_resume_pdf(path: Path) -> ResumeData:
    """Parse a text-based PDF resume into structured ResumeData.

    Uses pypdf's layout-aware extraction (``extraction_mode="layout"``) which
    preserves visual line separation far better than the default plain-text
    mode.  Text lines are then whitespace-normalised and run through the same
    heading heuristics as before.

    This does NOT work for scanned / image-only PDFs.
    """
    try:
        from pypdf import PdfReader
    except ImportError as e:
        raise ImportError(
            "pypdf is required to parse PDF resumes: pip install pypdf"
        ) from e

    reader = PdfReader(str(path))
    lines: list[str] = []
    for page in reader.pages:
        # Layout mode preserves visual line structure much better than plain mode
        text = page.extract_text(extraction_mode="layout") or ""
        lines.extend(text.splitlines())

    data = ResumeData()
    current_section: ResumeSection | None = None
    found_name = False

    for raw_line in lines:
        # Collapse multiple spaces (common in layout-mode output) to one space
        text = re.sub(r" {2,}", " ", raw_line).strip()
        if not text:
            continue

        if not found_name:
            data.name = text
            found_name = True
            continue

        if (
            not data.contact_info
            and current_section is None
            and not _is_pdf_heading(text)
        ):
            data.contact_info = text
            continue

        if _is_pdf_heading(text):
            current_section = ResumeSection(heading=text)
            data.sections.append(current_section)
        elif current_section is not None:
            current_section.items.append(text)
        else:
            if data.summary:
                data.summary += "\n" + text
            else:
                data.summary = text

    logger.info("Parsed resume (pdf): %s, %d sections", data.name, len(data.sections))
    return data


def _is_pdf_heading(text: str) -> bool:
    """Heuristic: identify section headings in PDF-extracted text lines.

    Checks (in order):
    1. Normalize whitespace; reject empty or long lines (> 60 chars).
    2. Reject lines that start with bullet markers (●, •, -, *).
    3. ALL-CAPS lines with > 2 characters are strong heading signals.
    4. Lines that equal, or start with, a known single- or multi-word
       section-heading keyword (followed by space, colon, or end-of-string).
    5. Lines that contain a known multi-word phrase as a substring (handles
       headings like "Familiar Programming Languages and Software").
    """
    # Normalize runs of whitespace to a single space
    text = re.sub(r"\s+", " ", text).strip()
    if not text or len(text) > 60:
        return False
    # Bullet / list markers are never section headings
    if re.match(r"^[●•\-\*]", text):
        return False
    # ALL-CAPS is a strong indicator (minimum 3 chars to avoid "ML", "AI", etc.)
    if text.isupper() and len(text) > 2:
        return True

    # Known single- and multi-word section heading keywords.
    # Detection rule: the normalised line *equals* the keyword OR *starts with*
    # the keyword followed by whitespace, a colon, or end-of-string.
    _heading_keywords = (
        # --- single-word headings ---
        "experience",
        "education",
        "skills",
        "summary",
        "objective",
        "projects",
        "certifications",
        "awards",
        "publications",
        "languages",
        "interests",
        "references",
        "employment",
        "career",
        "technologies",
        "tools",
        "coursework",
        "qualifications",
        "profile",
        "activities",
        "achievements",
        # --- multi-word headings (checked as exact / prefix matches) ---
        "work experience",
        "professional experience",
        "relevant experience",
        "work history",
        "technical skills",
        "key skills",
        "core skills",
        "relevant coursework",
        "professional summary",
        "career summary",
        "personal statement",
        "volunteer experience",
    )
    lower = text.lower()
    for kw in _heading_keywords:
        if lower == kw or re.match(rf"^{re.escape(kw)}[\s:,]", lower):
            return True

    # Substring match for multi-word phrases that may appear inside a longer
    # heading (e.g. "Familiar Programming Languages and Software").
    _phrase_substrings = (
        "programming languages",
        "languages and software",
    )
    return any(phrase in lower for phrase in _phrase_substrings)


def _parse_resume_md(path: Path) -> ResumeData:
    """Parse a Markdown resume into structured ResumeData.

    Sections are delimited by ATX headings (# / ## / ###).
    The first heading or first non-empty line is treated as the name.
    """
    content = path.read_text(encoding="utf-8")
    lines = content.splitlines()

    data = ResumeData()
    current_section: ResumeSection | None = None
    found_name = False

    for raw_line in lines:
        text = raw_line.strip()
        if not text:
            continue

        heading_match = re.match(r"^#{1,3}\s+(.*)", text)
        if heading_match:
            heading_text = heading_match.group(1).strip()
            if not found_name:
                data.name = heading_text
                found_name = True
            else:
                current_section = ResumeSection(heading=heading_text)
                data.sections.append(current_section)
            continue

        # Strip inline markdown (bold, italic, links, code)
        clean = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)  # links
        clean = re.sub(
            r"[*_`]{1,2}([^*_`]+)[*_`]{1,2}", r"\1", clean
        )  # bold/italic/code
        clean = re.sub(r"^[-*+]\s+", "", clean)  # list bullets

        if not found_name:
            data.name = clean
            found_name = True
            continue

        if not data.contact_info and current_section is None:
            data.contact_info = clean
            continue

        if current_section is not None:
            current_section.items.append(clean)
        else:
            if data.summary:
                data.summary += "\n" + clean
            else:
                data.summary = clean

    logger.info("Parsed resume (md): %s, %d sections", data.name, len(data.sections))
    return data


def generate_resume(data: ResumeData, template_path: Path, output_path: Path) -> Path:
    """Generate a tailored .docx resume from ResumeData using an existing .docx as template.

    Strategy: copy the template, then replace content while preserving formatting.
    For MVP, we create a clean document based on the template's styles.
    """
    # Copy the template to preserve styles and formatting
    template_doc = Document(str(template_path))
    doc = Document(str(template_path))

    # Clear all content from the copy
    for para in doc.paragraphs[:]:
        p_element = para._element
        p_element.getparent().remove(p_element)

    # Also clear tables if any
    for table in doc.tables[:]:
        t_element = table._element
        t_element.getparent().remove(t_element)

    # Get the heading style from the template
    heading_style = None
    normal_style = None
    for style in template_doc.styles:
        if style.name and "heading" in style.name.lower() and "1" in style.name:
            heading_style = style.name
        if style.name == "Normal":
            normal_style = style.name
    heading_style = heading_style or "Heading 1"
    normal_style = normal_style or "Normal"

    # Build the new document
    # Name
    name_para = doc.add_paragraph(data.name)
    _apply_safe_style(name_para, "Title", doc)

    # Contact info
    if data.contact_info:
        doc.add_paragraph(data.contact_info)

    # Summary
    if data.summary:
        doc.add_paragraph(data.summary)

    # Sections
    for section in data.sections:
        _apply_safe_style(doc.add_paragraph(section.heading), heading_style, doc)
        for item in section.items:
            # Use bullet style if available
            p = doc.add_paragraph(item)
            _apply_safe_style(p, "List Bullet", doc)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    logger.info("Generated resume at %s", output_path)
    return output_path


def _apply_safe_style(paragraph: Paragraph, style_name: str, doc: Document) -> None:
    """Apply a style if it exists, otherwise fall back to Normal."""
    try:
        paragraph.style = doc.styles[style_name]
    except KeyError:
        pass  # Keep default style
