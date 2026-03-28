"""Tests for resume parsing (PDF, docx, Markdown).

The PDF tests use the real CV_Resume.pdf file.  That file is in .gitignore
and must be present locally to run the PDF tests.  The tests are skipped
automatically when the file is not found.

Run with:
    pytest tests/test_resume.py -v
"""

from __future__ import annotations

import textwrap
from io import BytesIO
from pathlib import Path

import pytest
from docx import Document

from applybot.profile.resume import (
    ResumeData,
    ResumeSection,
    _is_pdf_heading,
    _parse_resume_md,
    _parse_resume_pdf,
    parse_resume,
)

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

_REAL_PDF = Path(__file__).parent.parent / "CV_Resume.pdf"
requires_real_pdf = pytest.mark.skipif(
    not _REAL_PDF.exists(),
    reason="CV_Resume.pdf not present (gitignored personal file)",
)


def _make_docx(paragraphs: list[dict[str, object]]) -> BytesIO:
    """Build a minimal .docx in memory from a list of paragraph descriptors.

    Each dict has 'text', optional 'style' (Word style name), optional
    'bold' (bool, applied to all runs when no named style is given).
    """
    doc = Document()
    for p in paragraphs:
        style = p.get("style")
        if style:
            para = doc.add_paragraph(str(p["text"]), style=str(style))
        else:
            para = doc.add_paragraph(str(p["text"]))
            if p.get("bold"):
                for run in para.runs:
                    run.bold = True
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ---------------------------------------------------------------------------
# _is_pdf_heading — unit tests
# ---------------------------------------------------------------------------


class TestIsPdfHeading:
    def test_all_caps_short_is_heading(self):
        assert _is_pdf_heading("EDUCATION") is True

    def test_all_caps_with_spaces_is_heading(self):
        assert _is_pdf_heading("WORK EXPERIENCE") is True

    def test_single_char_all_caps_not_heading(self):
        # Single/two-char tokens like "A" or "ML" should not be headings
        assert _is_pdf_heading("A") is False
        assert _is_pdf_heading("ML") is False

    def test_empty_string_not_heading(self):
        assert _is_pdf_heading("") is False

    def test_long_line_not_heading(self):
        long = "This is a very long line that clearly contains regular resume content and goes on and on"
        assert _is_pdf_heading(long) is False

    def test_bullet_line_not_heading(self):
        assert _is_pdf_heading("● Led a team to build ML pipelines") is False
        assert _is_pdf_heading("- Designed microservices architecture") is False
        assert _is_pdf_heading("• Created a new feature") is False

    # Single-word keyword headings
    def test_education_is_heading(self):
        assert _is_pdf_heading("Education") is True

    def test_experience_is_heading(self):
        assert _is_pdf_heading("Experience") is True

    def test_skills_is_heading(self):
        assert _is_pdf_heading("Skills") is True

    def test_projects_is_heading(self):
        assert _is_pdf_heading("Projects") is True

    def test_summary_is_heading(self):
        assert _is_pdf_heading("Summary") is True

    def test_certifications_is_heading(self):
        assert _is_pdf_heading("Certifications") is True

    # Multi-word keyword headings (key improvement over old code)
    def test_work_experience_is_heading(self):
        assert _is_pdf_heading("Work Experience") is True

    def test_professional_experience_is_heading(self):
        assert _is_pdf_heading("Professional Experience") is True

    def test_technical_skills_is_heading(self):
        assert _is_pdf_heading("Technical Skills") is True

    def test_programming_languages_is_heading(self):
        assert _is_pdf_heading("Familiar Programming Languages and Software") is True

    def test_relevant_coursework_is_heading(self):
        assert _is_pdf_heading("Relevant Coursework") is True

    def test_key_skills_is_heading(self):
        assert _is_pdf_heading("Key Skills") is True

    # Negative cases — content lines that should NOT be headings
    def test_job_title_not_heading(self):
        assert (
            _is_pdf_heading("Machine Learning Scientist at ASUS Robotics & AI Center")
            is False
        )

    def test_bullet_content_not_heading(self):
        assert (
            _is_pdf_heading(
                "● Led an engineering team to build a production ML pipeline"
            )
            is False
        )

    def test_sentence_with_keyword_not_heading(self):
        # A full sentence that happens to contain "experience" is not a heading
        assert _is_pdf_heading("5 years of experience in machine learning") is False

    def test_whitespace_normalized(self):
        # Extra spaces around text should not prevent detection
        assert _is_pdf_heading("  Education  ") is True
        assert _is_pdf_heading("Work  Experience") is True


# ---------------------------------------------------------------------------
# _parse_resume_md — unit tests with synthetic markdown
# ---------------------------------------------------------------------------


class TestParseResumeMd:
    def _parse_from_string(self, content: str, tmp_path: Path) -> ResumeData:
        p = tmp_path / "resume.md"
        p.write_text(textwrap.dedent(content), encoding="utf-8")
        return _parse_resume_md(p)

    def test_name_extracted_from_h1(self, tmp_path):
        data = self._parse_from_string(
            """
            # Jane Doe
            jane@example.com

            ## Experience
            - Senior Engineer at Acme
            """,
            tmp_path,
        )
        assert data.name == "Jane Doe"

    def test_contact_info_extracted(self, tmp_path):
        data = self._parse_from_string(
            """
            # Jane Doe
            jane@example.com | 555-1234

            ## Experience
            - Senior Engineer at Acme
            """,
            tmp_path,
        )
        assert "jane@example.com" in data.contact_info

    def test_sections_extracted(self, tmp_path):
        data = self._parse_from_string(
            """
            # Jane Doe
            jane@example.com

            ## Experience
            - Senior Engineer at Acme
            - Led a team of 5

            ## Education
            - B.Sc. Computer Science
            """,
            tmp_path,
        )
        headings = [s.heading for s in data.sections]
        assert "Experience" in headings
        assert "Education" in headings

    def test_section_items_extracted(self, tmp_path):
        data = self._parse_from_string(
            """
            # Jane Doe
            jane@example.com

            ## Skills
            - Python
            - Kubernetes
            """,
            tmp_path,
        )
        skills = data.get_section("Skills")
        assert skills is not None
        assert "Python" in skills.items
        assert "Kubernetes" in skills.items

    def test_summary_before_first_section(self, tmp_path):
        data = self._parse_from_string(
            """
            # Jane Doe
            jane@example.com
            Experienced ML engineer specializing in NLP.

            ## Experience
            - Lead at Acme
            """,
            tmp_path,
        )
        assert "ML engineer" in data.summary

    def test_inline_markdown_stripped(self, tmp_path):
        data = self._parse_from_string(
            """
            # **Jane Doe**
            jane@example.com

            ## Skills
            - **Python**, *Kubernetes*, `docker`
            """,
            tmp_path,
        )
        skills = data.get_section("Skills")
        assert skills is not None
        assert any("Python" in item for item in skills.items)

    def test_empty_file_returns_empty_data(self, tmp_path):
        data = self._parse_from_string("", tmp_path)
        assert data.name == ""
        assert data.sections == []

    def test_parse_resume_dispatches_md(self, tmp_path):
        p = tmp_path / "resume.md"
        p.write_text("# Test\n\n## Skills\n- Python\n", encoding="utf-8")
        data = parse_resume(p)
        assert data.name == "Test"

    def test_unsupported_extension_raises(self, tmp_path):
        p = tmp_path / "resume.txt"
        p.write_text("John Doe")
        with pytest.raises(ValueError, match="Unsupported"):
            parse_resume(p)


# ---------------------------------------------------------------------------
# _parse_resume_docx — unit tests with synthetic .docx
# ---------------------------------------------------------------------------


class TestParseResumeDocx:
    def _make_and_parse(
        self, paragraphs: list[dict[str, object]], tmp_path: Path
    ) -> ResumeData:
        buf = _make_docx(paragraphs)
        p = tmp_path / "resume.docx"
        p.write_bytes(buf.getvalue())
        return parse_resume(p)

    def test_name_is_first_paragraph(self, tmp_path):
        data = self._make_and_parse(
            [
                {"text": "John Smith"},
                {"text": "john@example.com"},
                {"text": "Experience", "style": "Heading 1"},
                {"text": "Engineer at Acme"},
            ],
            tmp_path,
        )
        assert data.name == "John Smith"

    def test_contact_info_second_paragraph(self, tmp_path):
        data = self._make_and_parse(
            [
                {"text": "John Smith"},
                {"text": "john@example.com | 555-1234"},
                {"text": "Experience", "style": "Heading 1"},
                {"text": "Engineer at Acme"},
            ],
            tmp_path,
        )
        assert "john@example.com" in data.contact_info

    def test_heading_style_starts_section(self, tmp_path):
        data = self._make_and_parse(
            [
                {"text": "John Smith"},
                {"text": "john@example.com"},
                {"text": "Experience", "style": "Heading 1"},
                {"text": "Engineer at Acme"},
                {"text": "Education", "style": "Heading 2"},
                {"text": "B.Sc. CS"},
            ],
            tmp_path,
        )
        headings = [s.heading for s in data.sections]
        assert "Experience" in headings
        assert "Education" in headings

    def test_bold_paragraph_starts_section(self, tmp_path):
        data = self._make_and_parse(
            [
                {"text": "John Smith"},
                {"text": "john@example.com"},
                {"text": "Skills", "bold": True},
                {"text": "Python, Kubernetes"},
            ],
            tmp_path,
        )
        assert data.get_section("Skills") is not None

    def test_section_items_collected(self, tmp_path):
        data = self._make_and_parse(
            [
                {"text": "John Smith"},
                {"text": "john@example.com"},
                {"text": "Skills", "style": "Heading 1"},
                {"text": "Python"},
                {"text": "Kubernetes"},
            ],
            tmp_path,
        )
        skills = data.get_section("Skills")
        assert skills is not None
        assert "Python" in skills.items
        assert "Kubernetes" in skills.items

    def test_pre_heading_text_is_summary(self, tmp_path):
        data = self._make_and_parse(
            [
                {"text": "John Smith"},
                {"text": "john@example.com"},
                {"text": "Experienced ML engineer."},
                {"text": "Skills", "style": "Heading 1"},
                {"text": "Python"},
            ],
            tmp_path,
        )
        assert "ML engineer" in data.summary


# ---------------------------------------------------------------------------
# _parse_resume_pdf — tests using the real CV_Resume.pdf
# ---------------------------------------------------------------------------


@requires_real_pdf
class TestParseRealPdf:
    @pytest.fixture(scope="class")
    def parsed(self) -> ResumeData:
        return _parse_resume_pdf(_REAL_PDF)

    def test_name_extracted(self, parsed: ResumeData):
        assert parsed.name != ""
        # The resume belongs to Sumanth Aluri
        assert "Sumanth" in parsed.name or "Aluri" in parsed.name

    def test_contact_info_contains_email(self, parsed: ResumeData):
        assert parsed.contact_info != ""
        assert "@" in parsed.contact_info

    def test_education_section_present(self, parsed: ResumeData):
        headings = [s.heading.lower() for s in parsed.sections]
        assert any(
            "education" in h for h in headings
        ), f"No education section found. Sections: {headings}"

    def test_experience_section_present(self, parsed: ResumeData):
        headings = [s.heading.lower() for s in parsed.sections]
        assert any(
            "experience" in h for h in headings
        ), f"No experience section found. Sections: {headings}"

    def test_skills_section_present(self, parsed: ResumeData):
        headings = [s.heading.lower() for s in parsed.sections]
        # Could be "Skills", "Technical Skills", "Programming Languages and Software", etc.
        assert any(
            "skill" in h or "technolog" in h or "language" in h for h in headings
        ), f"No skills section found. Sections: {headings}"

    def test_projects_section_present(self, parsed: ResumeData):
        headings = [s.heading.lower() for s in parsed.sections]
        assert any(
            "project" in h for h in headings
        ), f"No projects section found. Sections: {headings}"

    def test_at_least_four_sections(self, parsed: ResumeData):
        assert len(parsed.sections) >= 4, (
            f"Expected at least 4 sections, got {len(parsed.sections)}: "
            f"{[s.heading for s in parsed.sections]}"
        )

    def test_experience_section_has_items(self, parsed: ResumeData):
        exp = next(
            (s for s in parsed.sections if "experience" in s.heading.lower()), None
        )
        assert exp is not None
        assert len(exp.items) > 0, "Experience section has no items"

    def test_education_section_has_items(self, parsed: ResumeData):
        edu = next(
            (s for s in parsed.sections if "education" in s.heading.lower()), None
        )
        assert edu is not None
        assert len(edu.items) > 0, "Education section has no items"

    def test_no_duplicate_sections(self, parsed: ResumeData):
        headings = [s.heading.lower() for s in parsed.sections]
        assert len(headings) == len(
            set(headings)
        ), f"Duplicate section headings: {headings}"

    def test_parse_resume_dispatches_pdf(self):
        data = parse_resume(_REAL_PDF)
        assert data.name != ""


# ---------------------------------------------------------------------------
# ResumeData helpers
# ---------------------------------------------------------------------------


class TestResumeData:
    def test_get_section_case_insensitive(self):
        data = ResumeData(sections=[ResumeSection(heading="Skills", items=["Python"])])
        assert data.get_section("skills") is not None
        assert data.get_section("SKILLS") is not None

    def test_get_section_missing_returns_none(self):
        data = ResumeData()
        assert data.get_section("NonExistent") is None

    def test_round_trip_serialization(self):
        original = ResumeData(
            name="Jane Doe",
            contact_info="jane@example.com",
            summary="Experienced engineer.",
            sections=[
                ResumeSection(heading="Skills", items=["Python", "Kubernetes"]),
                ResumeSection(heading="Education", items=["B.Sc. CS"]),
            ],
        )
        restored = ResumeData.from_dict(original.to_dict())
        assert restored.name == original.name
        assert restored.contact_info == original.contact_info
        assert restored.summary == original.summary
        assert len(restored.sections) == len(original.sections)
        assert restored.sections[0].heading == "Skills"
        assert restored.sections[0].items == ["Python", "Kubernetes"]
